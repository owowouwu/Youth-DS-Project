# -*- coding: utf-8 -*-

import pandas as pd
from docx import Document

tsdr = Document('../raw_data/TSDR_2018-_Supplementary_Data_Report.docx')

# processes a table in the TSDR data set and returns it as a dictionary
def tsdrTable(n):
    data = {}

    for row in tsdr.tables[n].rows[1:len(tsdr.tables[n].rows)]:
        for i in range(0,len(row.cells), 2):
            if(row.cells[i].text == ''):
                break
            data[row.cells[i].text] = row.cells[i+1].text
    return data

# processes the required tables and puts everything into one dataframe

tableIndex = {'ecTeachers':59, 'ecStudents':64, 'psGovTeachers':124, 'psCTeachers':149, 'psStudents':158,
                   'ssGovTeachers':219, 'ssCTeachers':244, 'ssStudents':253}

df = pd.DataFrame([tsdrTable(i) for i in tableIndex.values()]).transpose()
df.columns = list(tableIndex.keys())
df = df.dropna()

# text processing
## replace all the less '<5' entries with 2.5
df = df.replace('<5', '2.5')
## make data numeric 
df = df.apply(lambda x: x.str.replace(',', ''), axis = 1).apply(pd.to_numeric)
df['tsRatio'] = (df['psGovTeachers'] + df['psCTeachers'] + df['ssGovTeachers'] +
                 df['ssCTeachers']) / (df['psStudents'] + df['ssStudents'])

df = df.sort_index()

df.to_csv('../wrangled/tsdr.csv')
