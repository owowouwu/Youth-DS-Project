# -*- coding: utf-8 -*-

"""
processes raw documents
"""

import pandas as pd
from docx import Document

# processing the aedc_raw excel file
aedc = pd.read_excel('./wrangled/AEDC_raw.xlsx')
aedc['LGA'] =  aedc['LGA'].replace('\([a-zA-Z.]*\)','', regex = True)
aedc = aedc.loc[:, ~aedc.columns.str.match('Unnamed')]
aedc = aedc.dropna()
aedc['LGA'] = aedc['LGA'].str.strip()
aedc = aedc.set_index('LGA')
# dropped here to remove incomplete data
aedc = aedc.drop('Queenscliffe')
aedc.to_csv('./wrangled/AEDC.csv')

# processing the depression_raw excel file
depression = pd.read_excel('./wrangled/depression_raw.xlsx', engine = 'openpyxl')
depression = depression.drop([0,1])
depression = depression[['Unnamed: 0','Persons']]
depression = depression.rename(columns={'Unnamed: 0': 'LGA', 'Persons': 'Depression Rate'})
depression['LGA'] = depression['LGA'].replace('\([a-zA-Z]*\)','', regex = True)
depression['LGA'] = depression['LGA'].str.strip()
depression = depression.set_index('LGA')
depression = depression.drop('Victoria')
depression.to_csv('./wrangled/depression.csv')

# processing teacher supply and demand report
tsdr = Document('../raw_data/TSDR_2018-_Supplementary_Data_Report.docx')

# processes a table in the TSDR data set and returns it as a dictionary
def tsdrTable(n):
    data = {}

    for row in tsdr.tables[n].rows[1:len(tsdr.tables[n].rows)]:
        for i in range(0,len(row.cells), 2):
            if(row.cells[i].text == ''):
                break
            data[row.cells[i].text] = row.cells[i+1].text
    return data

# processes the required tables and puts everything into one dataframe

tableIndex = {'ecTeachers':59, 'ecStudents':64, 'psGovTeachers':124, 'psCTeachers':149, 'psStudents':158,
                   'ssGovTeachers':219, 'ssCTeachers':244, 'ssStudents':253}

df = pd.DataFrame([tsdrTable(i) for i in tableIndex.values()]).transpose()
df.columns = list(tableIndex.keys())
df = df.dropna()

# text processing
## replace all the less '<5' entries with 2.5
df = df.replace('<5', '2.5')
## make data numeric 
df = df.apply(lambda x: x.str.replace(',', ''), axis = 1).apply(pd.to_numeric)
df['tsRatio'] = (df['psGovTeachers'] + df['ssGovTeachers'] +
                 df['ssCTeachers']) / (df['psStudents'] + df['ssStudents'])

df = df.sort_index()

df.to_csv('../wrangled/tsdr.csv')
df['tsRatio'].to_csv('../wrangled/tsratio.csv')